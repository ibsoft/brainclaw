import csv
import hashlib
import io
import json
import logging
import re
import uuid
from pathlib import Path
from typing import Any

from fastapi import HTTPException, UploadFile, status

from app.config import Settings
from app.db import Database
from app.embeddings import EmbeddingProvider
from app.faiss_store import FaissStore
from app.memory_service import chunk_content, redact_for_log, reject_obvious_secrets, row_to_memory
from app.schemas import FileSearchRequest, MemoryAddRequest


logger = logging.getLogger("brainclaw.files")

ALLOWED_EXTENSIONS = {
    ".txt",
    ".md",
    ".pdf",
    ".docx",
    ".csv",
    ".json",
    ".log",
    ".py",
    ".cs",
    ".js",
    ".yaml",
    ".yml",
    ".conf",
    ".service",
    ".html",
}

DANGEROUS_EXTENSIONS = {
    ".bat",
    ".cmd",
    ".com",
    ".dll",
    ".dmg",
    ".exe",
    ".jar",
    ".msi",
    ".ps1",
    ".scr",
    ".sh",
    ".vbs",
}


def sanitize_filename(filename: str) -> str:
    name = Path(filename or "upload").name
    return re.sub(r"[^A-Za-z0-9._ -]", "_", name).strip(" .") or "upload"


def parse_tags(raw_tags: str | None) -> list[str]:
    if not raw_tags:
        return []
    stripped = raw_tags.strip()
    if not stripped:
        return []
    try:
        parsed = json.loads(stripped)
        if isinstance(parsed, list):
            values = [str(item) for item in parsed]
        else:
            values = [str(parsed)]
    except json.JSONDecodeError:
        values = [item.strip() for item in stripped.split(",")]
    return MemoryAddRequest(
        agent_id="validator",
        workspace="validator",
        source="validator",
        memory_type="validator",
        content="validator",
        tags=values,
    ).tags


def decode_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="pypdf is required for PDF ingestion") from exc

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="python-docx is required for DOCX ingestion") from exc

    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            table_cells.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs + table_cells)


def extract_csv(data: bytes) -> str:
    text = decode_text(data)
    rows = []
    reader = csv.reader(io.StringIO(text))
    for row in reader:
        rows.append(" | ".join(row))
    return "\n".join(rows)


def extract_json(data: bytes) -> str:
    text = decode_text(data)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return text
    return json.dumps(parsed, indent=2, ensure_ascii=False)


def extract_text(data: bytes, extension: str) -> str:
    if extension == ".pdf":
        return extract_pdf(data)
    if extension == ".docx":
        return extract_docx(data)
    if extension == ".csv":
        return extract_csv(data)
    if extension == ".json":
        return extract_json(data)
    return decode_text(data)


class FileService:
    def __init__(self, settings: Settings, db: Database, embeddings: EmbeddingProvider, faiss_store: FaissStore):
        self.settings = settings
        self.db = db
        self.embeddings = embeddings
        self.faiss_store = faiss_store

    async def upload_file(self, agent_id: str, workspace: str, source: str, raw_tags: str | None, upload: UploadFile) -> dict[str, Any]:
        data = await upload.read(self.settings.max_upload_bytes + 1)
        return self.ingest_file_bytes(
            agent_id=agent_id,
            workspace=workspace,
            source=source,
            raw_tags=raw_tags,
            filename=upload.filename or "upload",
            content_type=upload.content_type,
            data=data,
        )

    def ingest_file_bytes(
        self,
        agent_id: str,
        workspace: str,
        source: str,
        raw_tags: str | None,
        filename: str,
        content_type: str | None,
        data: bytes,
    ) -> dict[str, Any]:
        agent_id = agent_id.strip()
        workspace = workspace.strip()
        source = source.strip()
        if not agent_id or not workspace or not source:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="agent_id, workspace, and source are required")

        original_filename = sanitize_filename(filename)
        extension = Path(original_filename).suffix.lower()
        if extension in DANGEROUS_EXTENSIONS or extension not in ALLOWED_EXTENSIONS:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="unsupported or dangerous file extension")

        size_bytes = len(data)
        if size_bytes == 0:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="empty files are not supported")
        if size_bytes > self.settings.max_upload_bytes:
            raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="file exceeds maximum upload size")

        sha256 = hashlib.sha256(data).hexdigest()
        duplicate = self.db.find_active_file_by_sha(sha256, agent_id, workspace)
        if duplicate is not None:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail={"message": "duplicate file", "file_id": int(duplicate["id"]), "sha256": sha256},
            )

        try:
            text = extract_text(data, extension).strip()
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="failed to extract text from file") from exc
        if not text:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="no extractable text found")
        reject_obvious_secrets(text)
        if len(text) > self.settings.max_content_chars:
            text = text[: self.settings.max_content_chars]

        stored_filename = f"{uuid.uuid4().hex}{extension}"
        stored_path = (self.settings.upload_dir / stored_filename).resolve()
        stored_path.relative_to(self.settings.upload_dir.resolve())
        stored_path.write_bytes(data)

        tags = parse_tags(raw_tags)
        chunks = chunk_content(text, self.settings.chunk_size_chars, self.settings.chunk_overlap_chars)
        file_id = self.db.create_file(
            {
                "agent_id": agent_id,
                "workspace": workspace,
                "source": source,
                "original_filename": original_filename,
                "stored_filename": stored_filename,
                "extension": extension,
                "content_type": content_type,
                "sha256": sha256,
                "size_bytes": size_bytes,
                "tags": tags,
            },
            chunks,
        )
        chunk_rows = self.db.get_chunks_for_file(file_id)
        vectors = self.embeddings.embed_texts([row["text"] for row in chunk_rows])
        mappings = [{"type": "file", "chunk_id": int(row["id"]), "file_id": file_id} for row in chunk_rows]
        self.faiss_store.add(vectors, mappings, agent_id, workspace)
        logger.info(
            "file_uploaded",
            extra={
                "file_id": file_id,
                "agent_id": agent_id,
                "workspace": workspace,
                "source": source,
                "original_filename": original_filename,
                "stored_filename": stored_filename,
                "extension": extension,
                "sha256": sha256,
                "size_bytes": size_bytes,
                "tags": tags,
                "chunks": len(chunks),
                "text_preview": redact_for_log(text[:80]),
            },
        )
        return {
            "ok": True,
            "file_id": file_id,
            "message": "file uploaded and indexed",
            "sha256": sha256,
            "stored_filename": stored_filename,
            "chunks": len(chunks),
            "duplicate": False,
        }

    def search_files(self, request: FileSearchRequest) -> list[dict[str, Any]]:
        vector = self.embeddings.embed_query(request.query)
        probe_limit = min(max(request.top_k * 10, request.top_k), max(self.faiss_store.vector_count, request.top_k))
        raw_results = self.faiss_store.search(vector, probe_limit, request.agent_id, request.workspace)
        results: list[dict[str, Any]] = []
        seen_chunks: set[int] = set()
        required_tags = set(request.tags or [])

        for raw in raw_results:
            if raw.get("type") != "file" or raw["score"] < request.min_score:
                continue
            row = self.db.get_file_chunk_with_file(raw["chunk_id"])
            if row is None:
                continue
            item = row_to_memory(row)
            if item["agent_id"] != request.agent_id or item["workspace"] != request.workspace:
                continue
            if request.file_id and item["id"] != request.file_id:
                continue
            if required_tags and not required_tags.issubset(set(item["tags"])):
                continue
            if item["file_chunk_id"] in seen_chunks:
                continue
            seen_chunks.add(item["file_chunk_id"])
            results.append(
                {
                    "file_id": item["id"],
                    "file_chunk_id": item["file_chunk_id"],
                    "score": raw["score"],
                    "agent_id": item["agent_id"],
                    "workspace": item["workspace"],
                    "source": item["source"],
                    "original_filename": item["original_filename"],
                    "extension": item["extension"],
                    "sha256": item["sha256"],
                    "size_bytes": item["size_bytes"],
                    "chunk_text": item["chunk_text"],
                    "tags": item["tags"],
                    "created_at": item["created_at"],
                    "updated_at": item["updated_at"],
                }
            )
            if len(results) >= request.top_k:
                break

        logger.info(
            "file_search",
            extra={
                "agent_id": request.agent_id,
                "workspace": request.workspace,
                "top_k": request.top_k,
                "min_score": request.min_score,
                "file_id": request.file_id,
                "tags": request.tags or [],
                "matches": len(results),
                "query_preview": redact_for_log(request.query),
            },
        )
        return results

    def delete_file(self, file_id: int, agent_id: str, workspace: str) -> dict[str, Any]:
        deleted = self.db.delete_file(file_id, agent_id, workspace)
        if not deleted:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="file not found")
        rebuilt = self.reindex_files_and_memories(agent_id, workspace)
        logger.info("file_deleted", extra={"file_id": file_id, "agent_id": agent_id, "workspace": workspace})
        return {"ok": True, "id": file_id, "message": "file deleted", "details": rebuilt["details"]}

    def reindex_files_and_memories(self, agent_id: str | None = None, workspace: str | None = None) -> dict[str, Any]:
        if self.faiss_store.settings.isolate_indexes and not (agent_id and workspace):
            total_vectors = 0
            scopes = self.db.list_scopes()
            for scope in scopes:
                rebuilt = self.reindex_files_and_memories(scope["agent_id"], scope["workspace"])
                total_vectors += int(rebuilt["details"]["vectors"])
            return {
                "ok": True,
                "id": None,
                "message": "isolated file and memory indexes rebuilt",
                "details": {"vectors": total_vectors, "scopes": len(scopes)},
            }
        if agent_id and workspace:
            memory_rows = self.db.iter_active_chunks_scoped(agent_id, workspace)
            file_rows = self.db.iter_active_file_chunks_scoped(agent_id, workspace)
        else:
            memory_rows = self.db.iter_active_chunks()
            file_rows = self.db.iter_active_file_chunks()
        texts = [row["chunk_text"] for row in memory_rows] + [row["chunk_text"] for row in file_rows]
        vectors = self.embeddings.embed_texts(texts)
        mappings: list[dict[str, Any]] = [
            {"type": "memory", "chunk_id": int(row["chunk_id"]), "memory_id": int(row["id"])}
            for row in memory_rows
        ]
        mappings.extend(
            {"type": "file", "chunk_id": int(row["file_chunk_id"]), "file_id": int(row["id"])}
            for row in file_rows
        )
        self.faiss_store.rebuild(vectors, mappings, agent_id, workspace)
        return {
            "ok": True,
            "id": None,
            "message": "file and memory indexes rebuilt",
            "details": {"vectors": len(mappings), "memory_chunks": len(memory_rows), "file_chunks": len(file_rows)},
        }
